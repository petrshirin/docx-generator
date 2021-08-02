import os

import docx
from typing import Union
import csv
import json
from time import sleep

try:
    with open('config.json', 'r', encoding='utf-8') as settings_file:
        settings = json.loads(settings_file.read())
except FileNotFoundError:
    print('''Не найден файл настроек config.json
Он должен находиться рядом с файлом main.exe (main.py)''')
    sleep(2)
    exit(-1)

try:
    DIR_PATH = settings['DIR_PATH']
    TEMPLATE_ROUTES = settings['TEMPLATE_ROUTES']
    NUMBER_MARK_TO_WORD = settings['NUMBER_MARK_TO_WORD']
    DATA_FILES = settings['DATA_FILES']
    CERTIFICATION_LIST_TEMPLATES = settings['CERTIFICATION_LIST_TEMPLATES']
    SERVICE_CHARACTERISTIC_TEMPLATES = settings['SERVICE_CHARACTERISTIC_TEMPLATES']
except KeyError:
    print('Неверно указан файл настроек config.json')
    sleep(2)
    exit(-1)


def create_docx_file_by_template(template_path: str) -> docx.Document:
    docx_file = docx.Document(template_path)
    return docx_file


def update_cell_value(self: docx.document.Document,
                      table_id: int,
                      row_id: int,
                      col_id: int,
                      value: Union[str, list],
                      method: int = 1) -> None:
    """

    :param self:
    :param table_id:
    :param row_id:
    :param col_id:
    :param value:
    :param method:
    1 - ADD IN CELL,
    2 - REPLACE ALL CELL VALUE,
    3 - REPLACE '_' SYMBOLS,
    4 - DELETE VALUE,
    5 - DELETE VALUE BY STR
    :return:
    """
    paragraph = self.tables[table_id].rows[row_id].cells[col_id].paragraphs[-1]
    if method == 1:
        r = paragraph.add_run(value)
        r.italic = True
    elif method == 2:
        self.tables[table_id].rows[row_id].cells[col_id].text = ""
        r = paragraph.add_run(value)
        r.italic = True
    elif method == 3:
        self.tables[table_id].rows[row_id].cells[col_id].text = self.tables[table_id].rows[row_id].cells[
            col_id].text.replace('_', '')
        r = paragraph.add_run(value)
        r.italic = True
    elif method == 4:
        self.tables[table_id].rows[row_id].cells[col_id].text = ""
    elif method == 5:
        self.tables[table_id].rows[row_id].cells[col_id].text = self.tables[table_id].rows[row_id].cells[
            col_id].text.replace(value, '')
    elif method == 6:
        i = 0
        for run in paragraph.runs:
            if run.text == '{}':
                run.text = value[i]
                run.italic = True
                i += 1


def show_template_cols():
    docx_file = docx.Document(TEMPLATE_ROUTES['послужные карты'][0])
    i = 0
    for table in docx_file.tables:
        j = 0
        for row in table.rows:
            k = 0
            for cell in row.cells:
                g = 0
                for par in cell.paragraphs:
                    print(i, j, k, g, [run.text for run in par.runs])
                    g += 1
                k += 1
            j += 1
        i += 1


def show_template_paragraphs():
    docx_file = docx.Document(SERVICE_CHARACTERISTIC_TEMPLATES['сержант']['3'])
    i = 0
    for par in docx_file.paragraphs:
        print(i, par.text)
        i += 1


def create_posluj_cards(template: str, data_file: str, dir_to_save: str):
    remove_data_from_dir(f'{DIR_PATH}{dir_to_save}\\Послужные карты')
    count = 0
    try:
        os.mkdir(f'{DIR_PATH}{dir_to_save}\\Послужные карты')
    except FileExistsError:
        pass
    with open(data_file, newline='', encoding='utf-8') as f:
        reader = csv.DictReader(f, delimiter=';')
        for row in reader:
            docx_file = create_docx_file_by_template(template)
            update_cell_value(docx_file, 0, 1, 0, row['ФИО'], 1)
            update_cell_value(docx_file, 0, 2, 0, [row['Дата рождения']], 6)
            update_cell_value(docx_file, 0, 3, 0, row['Место рождения'], 1)
            update_cell_value(docx_file, 0, 4, 0, 'русский', 5)
            update_cell_value(docx_file, 0, 5, 1,
                              f'ФГБОУ ВО "Иркутский государственный университет", {row["Год окончания учебы"]}', 1)
            update_cell_value(docx_file, 0, 7, 2, [row['Серия'], row['Номер'], row['Дата выдачи'], row['Кем выдан']], 6)
            update_cell_value(docx_file, 0, 12, 1, row['ИНН'], 1)
            update_cell_value(docx_file, 1, 0, 1, row['ВК'], 1)
            update_cell_value(docx_file, 1, 1, 1, row['Адрес регистрации'], 1)
            update_cell_value(docx_file, 1, 2, 1, row['Фактический адрес'], 1)
            update_cell_value(docx_file, 2, 0, 0, row['ФИО'], 1)
            update_cell_value(docx_file, 2, 2, 0, row['Дата рождения'], 1)
            update_cell_value(docx_file, 2, 7, 0, [row['Рост'],
                                                   row['Голова'],
                                                   row['Обувь'],
                                                   row['Противогаз'],
                                                   row['Оммуниция']], 6)
            update_cell_value(docx_file, 2, 8, 1, row['Фактический адрес'], 1)
            docx_file.save(DIR_PATH + f"{dir_to_save}\\Послужные карты\\Послужная карта {row['ФИО']}.docx")
            count += 1
    return count


def remove_data_from_dir(dir_to_remove):
    for walk_position in os.walk(dir_to_remove):
        for file in walk_position[2]:
            if '.docx' in file:
                os.remove(f'{dir_to_remove}\\{file}')


def create_help_cards(template: str, data_file: str, dir_to_save: str):
    remove_data_from_dir(f'{DIR_PATH}{dir_to_save}\\Справки')
    count = 0
    try:
        os.mkdir(f'{DIR_PATH}{dir_to_save}\\Справки')
    except FileExistsError:
        pass
    with open(data_file, newline='', encoding='utf-8') as f:
        reader = csv.DictReader(f, delimiter=';')
        for row in reader:
            docx_file = create_docx_file_by_template(template)
            values = [row['ФИО в дательном'], NUMBER_MARK_TO_WORD[row['ВТП']], NUMBER_MARK_TO_WORD[row['Итог']]]
            i = 0
            for paragraph in docx_file.paragraphs:
                if i > len(values):
                    break
                try:
                    for run in paragraph.runs:
                        if run.text.strip() == '{}':
                            run.text = values[i]
                            i += 1
                            run.underline = True
                except Exception as e:
                    print(values)
                    raise e
            docx_file.save(DIR_PATH + f"{dir_to_save}\\Справки\\Справка {row['ФИО']}.docx")
            count += 1
    return count


def create_service_characteristics(template: dict, data_file: str, dir_to_save: str):
    remove_data_from_dir(f'{DIR_PATH}{dir_to_save}\\Служебные характеристики')
    count = 0
    try:
        os.mkdir(f'{DIR_PATH}{dir_to_save}\\Служебные характеристики')
    except FileExistsError:
        pass
    with open(data_file, newline='', encoding='utf-8') as f:
        reader = csv.DictReader(f, delimiter=';')
        for row in reader:
            try:
                docx_file = create_docx_file_by_template(template[row['Итог']])
            except KeyError:
                continue
            values = [row['ФИО в родительном'], row['ФИО']]
            i = 0
            for paragraph in docx_file.paragraphs:
                if i > len(values):
                    break
                try:
                    for run in paragraph.runs:
                        if run.text.strip() == '{}':
                            run.text = values[i]
                            i += 1
                            run.underline = True
                except Exception as e:
                    print(values)
                    raise e
            docx_file.save(
                DIR_PATH + f"{dir_to_save}\\Служебные характеристики\\Служебная характеристика {row['ФИО']}.docx")
            count += 1
    return count


def create_certifications_lists(template: dict, data_file: str, dir_to_save: str):
    remove_data_from_dir(f'{DIR_PATH}{dir_to_save}\\Аттестационные листы')
    count = 0
    try:
        os.mkdir(f'{DIR_PATH}{dir_to_save}\\Аттестационные листы')
    except FileExistsError:
        pass
    with open(data_file, newline='', encoding='utf-8') as f:
        reader = csv.DictReader(f, delimiter=';')
        for row in reader:
            try:
                docx_file = create_docx_file_by_template(template[row['Итог']])
            except KeyError:
                continue
            values = [row['ФИО в дательном'],
                      row['Дата рождения'],
                      row['Место рождения'],
                      row['Факультет'],
                      row['Год окончания учебы'],
                      row['Гражданская специальность']]
            i = 0
            for paragraph in docx_file.paragraphs:
                if i > len(values):
                    break
                try:
                    for run in paragraph.runs:
                        if run.text.strip() == '{}':
                            run.text = values[i]
                            i += 1
                            run.underline = True
                except Exception as e:
                    print(values)
                    raise e
            docx_file.save(DIR_PATH + f"{dir_to_save}\\Аттестационные листы\\Аттестационный лист {row['ФИО']}.docx")
            count += 1
    return count


if __name__ == '__main__':
    generate_variant = 5
    count_generated_documents = 0
    try:
        generate_variant = int(input('''Что хотите сгенерировать?
1 - Послужные карты
2 - Справки
3 - Аттестационные листы
4 - Служебные характеристики
5 - Все документы
Вариант генерации документов: '''))
    except ValueError:
        print('Ошибка в формате')
        sleep(2)
        exit(-1)
    for key, value in DATA_FILES['сержант'].items():
        try:
            os.mkdir(f'{DIR_PATH}{key} взвод')
        except FileExistsError:
            pass
        if generate_variant == 1 or generate_variant == 5:
            count_generated_documents += create_posluj_cards(TEMPLATE_ROUTES['послужные карты']['сержант'], value,
                                                             f'{key} взвод')
        if generate_variant == 2 or generate_variant == 5:
            count_generated_documents += create_help_cards(TEMPLATE_ROUTES['справки']['сержант'], value, f'{key} взвод')
        if generate_variant == 3 or generate_variant == 5:
            count_generated_documents += create_certifications_lists(CERTIFICATION_LIST_TEMPLATES['сержант'], value,
                                                                     f'{key} взвод')
        if generate_variant == 4 or generate_variant == 5:
            count_generated_documents += create_service_characteristics(SERVICE_CHARACTERISTIC_TEMPLATES['сержант'],
                                                                        value, f'{key} взвод')

    for key, value in DATA_FILES['рядовой'].items():
        try:
            os.mkdir(f'{DIR_PATH}{key} взвод')
        except FileExistsError:
            pass
        if generate_variant == 1 or generate_variant == 5:
            count_generated_documents += create_posluj_cards(TEMPLATE_ROUTES['послужные карты']['рядовой'], value,
                                                             f'{key} взвод')
        if generate_variant == 2 or generate_variant == 5:
            count_generated_documents += create_help_cards(TEMPLATE_ROUTES['справки']['рядовой'], value, f'{key} взвод')
        if generate_variant == 3 or generate_variant == 5:
            count_generated_documents += create_certifications_lists(CERTIFICATION_LIST_TEMPLATES['сержант'], value,
                                                                     f'{key} взвод')
        if generate_variant == 4 or generate_variant == 5:
            count_generated_documents += create_service_characteristics(SERVICE_CHARACTERISTIC_TEMPLATES['сержант'],
                                                                        value, f'{key} взвод')

    print(f'Генерация документов завершена, сгенерированно {count_generated_documents} файлов')
    sleep(5)
